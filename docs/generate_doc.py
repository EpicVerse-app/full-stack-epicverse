from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ── Colour palette ────────────────────────────────────────────────────────────
PURPLE   = RGBColor(0x43, 0x19, 0x6B)   # deep purple — headings
GOLD     = RGBColor(0xC5, 0xA3, 0x58)   # gold — sub-headings
DARK     = RGBColor(0x1A, 0x1A, 0x2E)   # near-black — body text
WHITE    = RGBColor(0xFF, 0xFF, 0xFF)
TH_BG    = "431968"                      # table header bg (hex, no #)
ALT_BG   = "F3EEF9"                      # alternating row bg

# ── Helper — set cell background ──────────────────────────────────────────────
def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)

# ── Helper — styled table ─────────────────────────────────────────────────────
def add_table(headers, rows, col_widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header row
    hdr = t.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        set_cell_bg(cell, TH_BG)
        p  = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold      = True
        run.font.size = Pt(10)
        run.font.color.rgb = WHITE

    # Data rows
    for r_idx, row_data in enumerate(rows):
        row = t.rows[r_idx + 1]
        bg  = ALT_BG if r_idx % 2 == 1 else "FFFFFF"
        for c_idx, cell_text in enumerate(row_data):
            cell = row.cells[c_idx]
            set_cell_bg(cell, bg)
            p   = cell.paragraphs[0]
            run = p.add_run(str(cell_text))
            run.font.size = Pt(9.5)
            run.font.color.rgb = DARK

    # Column widths
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[i].width = Inches(w)

    doc.add_paragraph()   # spacer
    return t

# ── Helper — heading ──────────────────────────────────────────────────────────
def h1(text):
    p   = doc.add_paragraph()
    run = p.add_run(text.upper())
    run.bold           = True
    run.font.size      = Pt(15)
    run.font.color.rgb = PURPLE
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(4)

def h2(text):
    p   = doc.add_paragraph()
    run = p.add_run(text)
    run.bold           = True
    run.font.size      = Pt(12)
    run.font.color.rgb = GOLD
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(3)

# ── Helper — body paragraph ───────────────────────────────────────────────────
def body(text):
    p   = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size      = Pt(10.5)
    run.font.color.rgb = DARK
    p.paragraph_format.space_after = Pt(4)

# ── Helper — bullet ───────────────────────────────────────────────────────────
def bullet(text, bold_part=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_part and bold_part in text:
        before, after = text.split(bold_part, 1)
        if before:
            r = p.add_run(before)
            r.font.size = Pt(10.5)
            r.font.color.rgb = DARK
        r = p.add_run(bold_part)
        r.bold = True
        r.font.size = Pt(10.5)
        r.font.color.rgb = DARK
        if after:
            r = p.add_run(after)
            r.font.size = Pt(10.5)
            r.font.color.rgb = DARK
    else:
        run = p.add_run(text)
        run.font.size = Pt(10.5)
        run.font.color.rgb = DARK

# ══════════════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ══════════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("EpicVerse")
run.bold           = True
run.font.size      = Pt(28)
run.font.color.rgb = PURPLE

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Technical Overview Document")
run.bold           = True
run.font.size      = Pt(16)
run.font.color.rgb = GOLD

doc.add_paragraph()

meta = [
    ("Prepared by", "Kriyora Engineering Team"),
    ("Date",        "May 2026"),
    ("Version",     "1.0"),
    ("Status",      "Confidential — Internal Use Only"),
]
for label, value in meta:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p.add_run(f"{label}:  ")
    r1.bold = True
    r1.font.size = Pt(11)
    r1.font.color.rgb = DARK
    r2 = p.add_run(value)
    r2.font.size = Pt(11)
    r2.font.color.rgb = DARK

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# PART 1 — GAME MODE DATA
# ══════════════════════════════════════════════════════════════════════════════
h1("Part 1 — Game Mode Data")

# ── Overview ──────────────────────────────────────────────────────────────────
h2("Overview")
body(
    "EpicVerse validates card combos across 7 game modes. Each mode represents a "
    "chapter of the story and has its own set of valid characters and combo rules "
    "stored in the database. Players combine a character card with an attribute card "
    "and the AI checks whether the combination is valid, invalid, or excluded."
)

# ── Game Modes ────────────────────────────────────────────────────────────────
h2("Game Modes")
body("Each mode only accepts specific character cards. Characters outside a mode's list are rejected without querying the database.")

add_table(
    ["Mode", "Valid Character Card Numbers"],
    [
        ["Mode 1", "1, 2, 3, 5, 6, 7, 8, 9, 10, 23, 24"],
        ["Mode 2", "1, 2, 3, 5, 6, 8, 9, 10, 19, 24"],
        ["Mode 3", "1, 2, 3, 5, 11, 12, 15, 23"],
        ["Mode 4", "1, 2, 3, 4, 15, 17, 18, 21"],
        ["Mode 5", "2, 4, 11, 13, 14, 18, 20, 21"],
        ["Mode 6", "1, 3, 4, 11, 13, 14, 18, 20, 21, 22"],
        ["Mode 7", "1, 2, 3, 4, 5, 6, 13"],
    ],
    col_widths=[1.2, 4.5],
)

# ── Database Schema ────────────────────────────────────────────────────────────
h2("Database Schema — card_combos Table")
body("Every combo record in the PostgreSQL database stores the following fields:")

add_table(
    ["Column", "Type", "Description"],
    [
        ["id",                       "BIGSERIAL",   "Auto-incremented primary key"],
        ["gameplay_mode",            "TEXT",        "Game mode name"],
        ["character",                "TEXT",        "Character name"],
        ["character_card_number",    "INTEGER",     "Character card number"],
        ["attribute",                "TEXT",        "Attribute name"],
        ["attribute_card_no",        "INTEGER",     "Attribute card number (25+)"],
        ["final_segment",            "TEXT",        "Game segment reference"],
        ["final_status",             "TEXT",        "Valid / Invalid / Excluded"],
        ["revised_scholar_reason",   "TEXT",        "Lore-accurate explanation"],
        ["valmiki_reference_anchor", "TEXT",        "Source scripture reference"],
        ["kanda",                    "TEXT",        "Story chapter (Kanda)"],
        ["shloka",                   "TEXT",        "Verse reference"],
        ["explanation_summarized",   "TEXT",        "Short summary of the reason"],
        ["created_at",               "TIMESTAMPTZ", "Record creation timestamp"],
    ],
    col_widths=[2.2, 1.2, 3.3],
)

# ── Combo Validation Logic ────────────────────────────────────────────────────
h2("Combo Validation Logic")
body("When a user asks 'Is card X and card Y a combo?' the system follows these steps:")

bullet("Extracts two card numbers from the user's speech in any language or format", "Extracts")
bullet("Checks whether the character card belongs to the selected game mode", "Checks")
bullet("Queries the card_combos table in PostgreSQL for that combination", "Queries")
bullet("Returns one of three results — Valid, Invalid, or Excluded", "Returns")
bullet("Reads the lore explanation (revised_scholar_reason) from the database record", "Reads")

doc.add_paragraph()
add_table(
    ["Result", "Meaning"],
    [
        ["Valid",    "The combination is a recognised and active game combo"],
        ["Invalid",  "The combination does not form a valid combo"],
        ["Excluded", "Technically valid but excluded from current gameplay"],
    ],
    col_widths=[1.2, 5.0],
)

# ── Caching Strategy ──────────────────────────────────────────────────────────
h2("Data Caching Strategy")
body("The system uses three layers to ensure fast and reliable data access:")

bullet("Layer 1 — Redis Cache: Recently queried combos are stored in Redis with a time-based expiry for instant lookup", "Layer 1")
bullet("Layer 2 — PostgreSQL: If the result is not cached, the system queries the database directly", "Layer 2")
bullet("Layer 3 — Excel RAM Cache: If the database is unavailable, the system falls back to an in-memory copy of the Excel data loaded at server startup", "Layer 3")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# PART 2 — TECHNOLOGY STACK
# ══════════════════════════════════════════════════════════════════════════════
h1("Part 2 — Technology Stack")

# ── Frontend ──────────────────────────────────────────────────────────────────
h2("Frontend — Mobile Application")
body(
    "The mobile app is built using Flutter (Dart), Google's cross-platform framework, "
    "allowing a single codebase to run on both Android and iOS. State management is "
    "handled by Riverpod."
)

add_table(
    ["Package", "Version", "Purpose"],
    [
        ["flutter",               "SDK",      "Core mobile framework (Android + iOS)"],
        ["flutter_riverpod",      "^3.3.1",   "Reactive state management"],
        ["dio",                   "^5.9.2",   "HTTP client for REST API calls"],
        ["web_socket_channel",    "^3.0.3",   "WebSocket connection to backend"],
        ["record",                "^6.2.0",   "Mic recording — PCM16 24kHz mono"],
        ["flutter_pcm_sound",     "^3.3.3",   "AI voice playback — PCM 24kHz"],
        ["audioplayers",          "^6.6.0",   "General audio playback"],
        ["speech_to_text",        "^7.0.0",   "Local speech recognition support"],
        ["firebase_core",         "^4.6.0",   "Firebase initialisation"],
        ["firebase_auth",         "^6.3.0",   "User authentication"],
        ["permission_handler",    "^12.0.1",  "Mic and storage permissions"],
        ["shared_preferences",    "^2.5.4",   "Local session storage"],
        ["uuid",                  "^4.3.3",   "Unique session ID generation"],
        ["image_picker",          "^1.2.1",   "Profile picture selection"],
        ["google_fonts",          "^8.0.2",   "Typography"],
        ["animations",            "^2.1.1",   "Screen transition animations"],
        ["intl",                  "^0.20.2",  "Internationalisation support"],
        ["flutter_launcher_icons","^0.14.4",  "App icon generation"],
    ],
    col_widths=[2.0, 1.0, 3.7],
)

# ── Audio Format ──────────────────────────────────────────────────────────────
h2("Audio Format")
add_table(
    ["Direction", "Format", "Sample Rate", "Channels"],
    [
        ["Mic → Backend",    "PCM 16-bit signed", "16 kHz", "Mono"],
        ["Backend → Speaker","PCM 16-bit signed", "24 kHz", "Mono"],
    ],
    col_widths=[1.8, 2.0, 1.4, 1.2],
)

# ── Backend ───────────────────────────────────────────────────────────────────
h2("Backend — Server Application")
body(
    "The backend is a Python 3.11 application built with FastAPI, running on Uvicorn "
    "(ASGI server). It is hosted on Google Cloud Run in the Mumbai region (asia-south1) "
    "for low latency. The server bridges the Flutter app and the OpenAI Realtime API "
    "over a persistent WebSocket connection."
)

add_table(
    ["Package", "Version", "Purpose"],
    [
        ["fastapi",               "0.103.2",  "Web framework — REST API and WebSocket"],
        ["uvicorn",               "0.23.2",   "ASGI server"],
        ["websockets",            "12.0",     "WebSocket client for OpenAI Realtime bridge"],
        ["asyncpg",               "0.29.0",   "Async PostgreSQL database driver"],
        ["pydantic-settings",     ">=2.0.0",  "Settings and environment variable validation"],
        ["openai",                ">=1.3.5",  "OpenAI API client"],
        ["firebase-admin",        "6.2.0",    "Firebase ID token verification"],
        ["redis",                 ">=5.0.0",  "Redis cache client"],
        ["numpy",                 ">=1.26.0", "PCM audio resampling (16kHz to 24kHz)"],
        ["pandas",                "2.1.1",    "Excel game data loading"],
        ["openpyxl",              "3.1.2",    "Excel file parsing"],
        ["python-multipart",      "0.0.6",    "File upload handling"],
        ["httpx",                 ">=0.24.0", "Async HTTP client"],
        ["google-cloud-storage",  "2.11.0",   "Google Cloud Storage file operations"],
        ["google-cloud-logging",  "3.6.0",    "Structured cloud logging"],
        ["python-dotenv",         "1.0.0",    "Local environment variable loading"],
    ],
    col_widths=[2.0, 1.1, 3.6],
)

# ── AI & Voice ────────────────────────────────────────────────────────────────
h2("AI and Voice Services")
body(
    "The entire voice experience runs through the OpenAI Realtime API. A single "
    "WebSocket connection handles speech recognition, language understanding, tool "
    "execution, and voice response — with no separate services to coordinate."
)

add_table(
    ["Service", "Provider", "Detail"],
    [
        ["Large Language Model (LLM)", "OpenAI GPT-4o Realtime",    "Language understanding, combo logic, tool calling"],
        ["Speech-to-Text (STT)",       "OpenAI Whisper-1",          "Auto-detects and transcribes 99 languages"],
        ["Text-to-Speech (TTS)",       "OpenAI Realtime — alloy",   "Streams PCM audio response to the user"],
        ["Translation",                "GPT-4o native",             "Translates responses into user's language (100+ languages)"],
    ],
    col_widths=[2.1, 2.1, 2.5],
)

# ── Infrastructure ────────────────────────────────────────────────────────────
h2("Cloud Infrastructure")

add_table(
    ["Service", "Provider", "Purpose"],
    [
        ["App Hosting",       "Google Cloud Run",        "Serverless backend — scales 1 to 20 instances automatically"],
        ["Database",          "Google Cloud SQL",        "PostgreSQL — primary game data and user records"],
        ["Cache",             "Redis",                   "Fast in-memory combo lookup cache"],
        ["Container Registry","Google Artifact Registry","Stores Docker images for deployment"],
        ["CI/CD Pipeline",    "Google Cloud Build",      "Automated build and deploy on every code release"],
        ["Secrets",           "Google Secret Manager",   "Stores API keys and database credentials securely"],
        ["Authentication",    "Firebase Authentication", "User login — Google Sign-In and email"],
        ["Email",             "SendGrid",                "Delivers OTP emails for account verification"],
        ["Source Control",    "GitHub",                  "Code repository and version history"],
    ],
    col_widths=[1.8, 2.0, 3.0],
)

# ── Cloud Run Config ──────────────────────────────────────────────────────────
h2("Cloud Run Configuration")

add_table(
    ["Parameter", "Value"],
    [
        ["Region",           "asia-south1 (Mumbai, India)"],
        ["Min Instances",    "1 — always warm, no cold start"],
        ["Max Instances",    "20 — auto-scales under load"],
        ["Memory",           "2 GB per instance"],
        ["vCPUs",            "2 per instance"],
        ["Max Concurrency",  "80 simultaneous connections per instance"],
    ],
    col_widths=[2.2, 4.5],
)

# ── Endpoint ──────────────────────────────────────────────────────────────────
h2("Production Endpoint")
body("The live backend is accessible at:")
p = doc.add_paragraph()
run = p.add_run("https://epicverse-backend-721191424605.asia-south1.run.app")
run.bold = True
run.font.size = Pt(10.5)
run.font.color.rgb = PURPLE

doc.add_paragraph()
doc.add_paragraph()

# ── Footer ────────────────────────────────────────────────────────────────────
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Kriyora Concepts Private Limited  |  Confidential  |  support@kriyora.com")
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

# ── Save ──────────────────────────────────────────────────────────────────────
out = r"e:\kriyora\EpicVerse\docs\EpicVerse-Technical-Overview.docx"
doc.save(out)
print(f"Saved: {out}")
